"""
Generate The Care Ranch Brochure v8 (improved) as .docx

Uses the original v7 .docx as a template so all fonts, colors, heading styles,
and theme definitions are inherited exactly. Only the content is replaced.

Applies all feedback from v7 review:
- Restructured sections (Who This Is For + Team moved earlier)
- Consistent poetic-yet-grounded voice throughout
- Consistent first-person plural POV
- Fixed Dutch-to-English translation points
- EMDR integrated naturally, EBAS acronym moved to Scott's bio
- Theory U reference removed
- Non-profit section bridged properly
- Standardized team bios
- Stronger CTA, closing quote repositioned
"""

from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
TEMPLATE = os.path.join(PROJECT_DIR, 'client-input', 'brochure-v7-original.docx')
OUTPUT_DIR = os.path.join(PROJECT_DIR, 'output')

# ---------------------------------------------------------------------------
# Load template and clear content
# ---------------------------------------------------------------------------

doc = Document(TEMPLATE)

# Remove all paragraphs and tables from the body, but keep section properties
body = doc.element.body
for child in list(body):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag in ('p', 'tbl'):
        body.remove(child)

# ---------------------------------------------------------------------------
# Helpers — match the original's formatting exactly
# ---------------------------------------------------------------------------

# The original (Google Docs export) stores formatting as paragraph-level and
# run-level overrides, not just in style definitions. We must replicate these:
#
#   Heading 1: run 23pt bold, para space_before=24pt
#   Heading 2: run 17pt bold, para space_after=4pt (space_before from style=18pt)
#   Heading 3: run 13pt bold, para space_before=14pt
#   Normal body: para space_after=12pt (majority of paragraphs)
#   Method/benefit items: bold prefix + \n + body, all in one 'normal' paragraph
#   No bullet styles — list items are plain 'normal' paragraphs

def h1(text):
    """Main title (Heading 1 — 23pt bold, 24pt space before)."""
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(24)
    for run in p.runs:
        run.font.size = Pt(23)
        run.font.bold = True
    return p

def h2(text):
    """Section heading (Heading 2 — 17pt bold, 4pt space after)."""
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(17)
        run.font.bold = True
    return p

def h3(text):
    """Sub-section heading (Heading 3 — 13pt bold, 14pt space before)."""
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(14)
    for run in p.runs:
        run.font.size = Pt(13)
        run.font.bold = True
    return p

def body(text, bold=False, italic=False, space_after=Pt(12)):
    """Normal paragraph with 12pt space after (default)."""
    p = doc.add_paragraph(style='normal')
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def body_mixed(parts, space_after=Pt(12)):
    """Normal paragraph with mixed formatting.
    parts: list of (text, bold, italic) tuples.
    """
    p = doc.add_paragraph(style='normal')
    p.paragraph_format.space_after = space_after
    for text, bld, ital in parts:
        run = p.add_run(text)
        if bld:
            run.bold = True
        if ital:
            run.italic = True
    return p

def item(text):
    """List-like item — plain paragraph (matching original, no bullet style)."""
    return body(text)

def spacer():
    """Empty paragraph as visual break (no extra spacing)."""
    p = doc.add_paragraph(style='normal')
    return p


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

# ===== TITLE =====
h1('THE CARE RANCH')

body(
    'An immersive leadership experience grounded in neuroscience '
    'and embodied intelligence — resulting in aligned leadership '
    'with clean vision, clarity of direction, and impactful outcomes.',
    italic=True)

# ===== THE CARE RANCH IN ARIZONA =====
h2('The Care Ranch in Arizona')

body('Set in the high Sonoran Desert of Arizona, The Care Ranch '
     'collaborates with the historic Tubac Ranch — built in the 1800s — '
     'preserving its essence and partnership with nature.')

body('The ranch is restored with care, respect, and restraint, '
     'honoring the landscape and its historic charm.')

body('Nothing is excessive or distracting from the transformative journey.')

body('The land is open. The rhythm is slow. The silence is your peace.',
     italic=True)

body('Here, simplicity calms the nervous system while the absence of excess '
     'reduces cognitive load. This historic space slows perception and '
     'grounds the body. Nothing here is accidental.')

body('Food is prepared with care, using whole seasonal ingredients that '
     'support the nervous system and reflect the rhythm of the land.')

body('Nature is not a backdrop — it is an active force.')

body('It regulates. It slows perception. It brings back what is essential.')

body('The desert does not add anything. It removes.', italic=True)

body('And in that space, vision becomes visible again.')

# ===== CLEAN VISION =====
h2('This Is the Place for Clean Vision')

body('The Care Ranch offers immersive leadership experiences designed '
     'for executives, senior leaders, and individuals who understand that '
     'clarity, resilience, and authority do not come from speed or control '
     '— but from being in alignment.')

body('Leading well in complexity requires a different kind of skill '
     '— the ability to restore congruence in the nervous system, '
     'engage emotional resilience, and embody vision for sustainable, '
     'grounded, and trustworthy leadership.')

# ===== WHO THIS IS FOR (moved up — per feedback) =====
h2('Who This Is For')

body('This experience is for leaders and individuals who sense:')

item('that something in their leadership or life is ready for recalibration,')
item('that there is a deeper layer to how they decide and lead,')
item('that they are ready to evolve into a new phase of development,')
item('and who want to reconnect with direction, clarity, and themselves.')

body('Whether the question is fully formed or still emerging, '
     'what participants share is a sense that something is ready to shift.')

# ===== THE CARE RANCH METHOD =====
h2('THE CARE RANCH METHOD')

body('Grounded in neuroscience and refined through practice, we reveal '
     'great leadership by cultivating a relationship with the body '
     'and its nervous system.')

body('How you lead and decide is shaped by how your system operates '
     '— especially under pressure.')

body('We develop leadership by restoring alignment in the body and '
     'nervous system, allowing for clear perception and more precise '
     'decision-making.')

body('This is one integrated, multilayered approach — working through '
     'multiple entry points.')

# Method components — bold title + \n + description, matching original format
body_mixed([
    ('Movement and Body Alignment\n', True, False),
    ('Working with posture, coordination, and movement patterns '
     'to reorganize neuromuscular pathways '
     'and restore alignment between body and brain.', False, False),
])

body_mixed([
    ('Equine Facilitated Learning\n', True, False),
    ('Horses respond to what is real — not what is intended. '
     'They reflect incongruence in timing, presence, and boundaries instantly, '
     'making underlying patterns visible without interpretation.', False, False),
])

body_mixed([
    ('Neuro-Kinesiology\n', True, False),
    ('Using muscle biofeedback to identify neurological stress patterns '
     'and bring unconscious responses into awareness. From there, '
     'the system reorganizes.', False, False),
])

body_mixed([
    ('Coaching and EMDR\n', True, False),
    ('Individual one-on-one coaching designed for reflection and '
     'integration, supporting the processing of underlying patterns and '
     'translating insight into concrete decisions and action. '
     'Where relevant, EMDR is used to help shift deeply rooted patterns, '
     'allowing new perspectives and responses to emerge.', False, False),
])

body_mixed([
    ('Nature and Physiology\n', True, False),
    ('The environment, the rhythm of the day, and the nourishment '
     'support regulation, recovery, and integration at a physiological level.', False, False),
])

body('Together, this creates a shift:')
body('From thinking to sensing.\n'
     'From reacting to responding.\n'
     'From fragmentation to alignment.', italic=True)

body('At the center lies what we call the Wheel of Congruence — '
     'where vision, emotion, body, and action come back into alignment.')

body('You learn to read your body and emotions as information. '
     'Not something to push through, but something that informs direction.')

# ===== OUR TEAM (moved up — per feedback: trust-builder) =====
h2('OUR TEAM')

body('A multidisciplinary team working at the intersection of neuroscience, '
     'embodiment, nature, and leadership — led by a core team and supported '
     'by an extended network of practitioners and specialists.')

# --- Margreet (Founder first — per feedback: fuller bio) ---
h3('Margreet Malenstein — Founder and Vision Holder, Lead Trainer')

body('As former co-owner of Bergman Clinics, an international healthcare '
     'group, I combine firsthand executive experience with a deep focus '
     'on the body, the nervous system, and preventive care.')

body('My work is grounded in a thorough understanding of human behavior '
     'and in how safety, alignment, and environment influence both healing '
     'and performance. Having navigated the pressures of executive '
     'leadership myself, I know what it means when the system is out of '
     'alignment — and what becomes possible when it is restored.')

body('Through horses and neuro-kinesiology, I support leaders in '
     'recognizing where alignment is present and where it is not — '
     'and in restoring congruence from within.')

body('The Care Ranch was born from a conviction that the principles '
     'we apply to leadership development should be accessible more broadly '
     '— which is why we also extend this work to young people '
     'through our non-profit branch.')

# --- Scott ---
h3('Scott Putman — Co-Creator, Method Architect and Lead Trainer')

body('My work focuses on the relationship between movement, '
     'the nervous system, and perception.')

body('With an international academic background (MFA, University of '
     'California), I developed EBAS — the Elemental Body Alignment '
     'System — as a way to work directly with the reorganization of '
     'neuromuscular pathways. Through movement, I support the body and '
     'brain in coming back into alignment, directly influencing how '
     'the nervous system regulates and how clarity is experienced.')

body('The body reflects the internal state. Through movement and '
     'Equine Facilitated Learning, that state becomes visible — '
     'and adjustable through direct feedback.')

body('Alignment is not explained. It is experienced. '
     'From there, the system reorganizes.')

# --- Lotte ---
h3('Lotte van Dam — Co-Creator, Method Architect and Lead Trainer')

body('My work focuses on personal leadership in complexity.')

body('With a background in Psychology and International Business '
     'Administration, and over 18 years of leadership experience across '
     'hospitality, design, and healthcare — including executive and '
     'board-level roles — I understand both the outer complexity of '
     'leadership and the inner landscape that shapes it.')

body('Rooted in neuroscience, neuroplasticity, and systemic coaching, '
     'I work at the intersection of perception, behavior, and the '
     'nervous system, using the body as a source of information.')

body('Through coaching, EMDR, and integration work, I support leaders '
     'in recognizing patterns, processing underlying responses, and '
     'restoring alignment between body, emotion, and action.')

body('This is where clarity emerges, congruence becomes visible, '
     'and leadership becomes grounded — with impact that extends '
     'into teams, organizations, and systems.')

# ===== THE EXPERIENCE =====
h2('THE EXPERIENCE')

body('The experience unfolds over six days — '
     'a rhythm of activation and integration. Structured, yet spacious.')

body('Mornings begin in the body: movement, alignment, regulation. '
     'During the day, the work deepens — equine sessions, '
     'neuro-kinesiology, and coaching. '
     'Always alternating between experience, reflection, and integration.')

body('Late afternoons create space for reflection and sauna house '
     'experiences. Evenings settle under a sky filled with stars.')

body('Each participant arrives with their own question. Sometimes clear. '
     'Sometimes not yet fully formed.')

body('Private sessions create precision. '
     'The group — intentionally small — creates reflection and perspective.',
     italic=True)

# ===== THE JOURNEY =====
h2('THE JOURNEY')

h3('Preparing for the Experience')

body('The journey begins before arrival. '
     'Each participant starts with an intake questionnaire and session, '
     'exploring context, key questions, and underlying patterns.')

body('A preparation guide supports participants in:')

item('Clarifying intention')
item('Increasing awareness')
item('Slowing down before entering the program')

body('This ensures participants arrive focused and connected '
     'to what is ready to be explored.')

h3('Six-Day Immersive Leadership Experience')

body('Participants enter an environment designed to regulate '
     'the nervous system. As the system settles:')

item('Perception sharpens')
item('Internal noise reduces')
item('Patterns become visible')

body('Through Equine Facilitated Learning, neuro-kinesiology, '
     'body alignment, coaching, and nature-based practices, '
     'insight is not only understood — it is experienced and embodied.')

h3('Integration — Sustaining Change')

body('After the immersion, the process continues.')

body('Participants are supported in translating insights into daily life '
     'and professional context through a structured integration phase, '
     'including follow-up sessions and integration guidance.')

body('The focus is on:')

item('Stabilizing new patterns')
item('Maintaining clarity under pressure')
item('Embedding change over time')

body('Because transformation is not defined by the experience itself '
     '— but by what remains.', bold=True)

# ===== WHAT LEADERS GAIN =====
h2('What Leaders Gain')

body_mixed([
    ('Clarity of purpose\n', True, False),
    ('A reconnection with what truly matters — beyond pressure, '
     'expectation, or noise. Not driven by thought alone, '
     'but aligned with the whole system.', False, False),
])

body_mixed([
    ('The ability to regulate under pressure\n', True, False),
    ('Learning how to access calm, focus, and precision in high-stakes '
     'environments. Not by pushing harder, but by regulating '
     'the nervous system.', False, False),
])

body_mixed([
    ('Embodied self-awareness\n', True, False),
    ('Deep insight into patterns, roles, and behaviors — '
     'and how they show up in relation to others. '
     'Felt, not just understood.', False, False),
])

body_mixed([
    ('Congruent leadership\n', True, False),
    ('When what you think, feel, and do align, decisions become cleaner, '
     'communication more direct, and leadership more effective.', False, False),
])

body_mixed([
    ('Sustainable change\n', True, False),
    ('Through integration and one-on-one guidance, insights are '
     'translated into daily life and leadership practice. '
     'This is not a peak experience — it is a shift that holds.', False, False),
])

body_mixed([
    ('Space to pause and reset\n', True, False),
    ('A rare interruption of speed. A place where slowing down '
     'is not a luxury, but the condition for clarity.', False, False),
])

body_mixed([
    ('A deeper connection to self and others\n', True, False),
    ('Through nature, horses, and human interaction, leaders experience '
     'what it means to be fully present — '
     'and to lead from that place.', False, False),
])

# ===== WHY THIS MATTERS FOR ORGANISATIONS (clearly separated) =====
h2('Why This Matters for Organizations')

body('For HR leaders and organizations investing in senior talent, '
     'this experience supports what traditional leadership programs '
     'often cannot reach:')

item('Preventive leadership care — addressing patterns before they become burnout')
item('Stronger leadership presence and trust within teams')
item('Sustainable performance under complexity')
item('A measurable shift in how leaders regulate, decide, and communicate')

body('Leaders return more coherent — not merely rested.', bold=True)

body('The impact extends beyond the individual leader into teams, '
     'culture, and long-term organizational health.')

# ===== FROM LEADERSHIP TO CONTRIBUTION (with bridge) =====
h2('From Leadership to Contribution')

body('The work at the heart of The Care Ranch — learning to regulate, '
     'to reconnect, and to find direction from within — '
     'is not only relevant for senior leaders. '
     'It is just as essential for young people who feel stuck, '
     'overwhelmed, or without direction.')

body('The Care Ranch operates as a hybrid model. '
     'The for-profit leadership programs fund the non-profit branch, '
     'making this work accessible for young individuals who would '
     'otherwise not have access.')

body('Through participation, you contribute to creating impact '
     'beyond the individual journey.')

body('Because learning to regulate, to reconnect, and to find direction '
     'in yourself should be available earlier — '
     'not only when things break down.', italic=True)

# ===== PRACTICAL INFORMATION =====
h2('PRACTICAL INFORMATION')

body_mixed([
    ('Location\n', True, False),
    ('The Care Ranch — Tubac Ranch, Arizona, USA', False, False),
])

body_mixed([
    ('Duration\n', True, False),
    ('6 days / 5 nights', False, False),
])

body_mixed([
    ('Group size\n', True, False),
    ('8–12 participants', False, False),
])

body_mixed([
    ('Includes\n', True, False),
    ('Full program and personal guidance\n'
     'Private sessions (coaching, kinesiology, equine work)\n'
     'Daily body alignment movement practice\n'
     'Accommodation (private casitas)\n'
     'Nourishing meals and non-alcoholic beverages\n'
     'On-site facilities (swimming pools, sauna house experience)\n'
     'Transfer to and from Tucson airport', False, False),
])

body_mixed([
    ('Not included\n', True, False),
    ('Flights to Tucson airport', False, False),
])

# ===== BEGIN YOUR JOURNEY =====
h2('Begin Your Journey')

# Closing quote BEFORE the CTA (per feedback)
body('"A place where heritage is honored, simplicity restores '
     'the nervous system, and leadership returns to congruence."',
     bold=True)

spacer()

body('Program dates, pricing, and availability are listed on the website.')

body('Participation is limited and based on application.')

body_mixed([
    ('Connect with us: ', True, False),
    ('www.thecareranch.com', False, False),
])

body('Where needed, we are happy to connect briefly to explore alignment '
     'and answer any questions.')


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

os.makedirs(OUTPUT_DIR, exist_ok=True)
output_path = os.path.join(OUTPUT_DIR, 'brochure-v8-improved.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
